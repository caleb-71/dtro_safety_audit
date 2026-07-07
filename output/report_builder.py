# output/report_builder.py
# 워드 보고서 자동 생성 모듈
# 표 가운데 정렬 + 컬럼 폭 자동 조절 적용

import logging
from datetime import datetime
from pathlib import Path

import matplotlib
import matplotlib.pyplot as plt
import matplotlib.font_manager as fm
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from config.settings import OUTPUT_DIR
from core.action_analyzer import (
    add_action_status, action_summary,
    action_rate_by_group, consulting_targets,
)

logger = logging.getLogger(__name__)


# ─────────────────────────────────────────
# AI 서술 문단 생성 (신규, 2026-07)
# ─────────────────────────────────────────
def _generate_narrative(topic: str, stats_text: str) -> str:
    """
    집계된 통계 수치를 근거로 보고서용 분석 서술 문단을 LLM으로 생성합니다.

    설계 원칙:
    - 통계에 있는 수치만 인용하도록 프롬프트에 명시 (환각 억제)
    - Ollama 미가동 등 오류 시 빈 문자열 반환 → 보고서 생성 자체는 계속됨
      (AI 서술은 '있으면 좋은' 부가 요소이므로 실패가 전체를 막으면 안 됨)
    """
    try:
        from core.llm_client import llm_chat
        from config.settings import OLLAMA_MODEL

        prompt = f"""당신은 대구교통공사 자체종합안전심사 결과보고서를 작성하는 안전관리 전문가입니다.
아래 통계 데이터를 근거로 '{topic}'에 대한 분석 문단을 작성하세요.

[통계 데이터]
{stats_text}

[작성 지침]
1. 통계 데이터에 있는 수치만 정확히 인용하세요 (추측·과장 금지)
2. 보고서 서술체로 3~5문장의 하나의 문단으로 작성하세요
3. 수치 인용 후 안전관리 관점의 시사점을 1가지 포함하세요
4. 제목이나 머리기호 없이 본문 문단만 작성하세요
5. 한국어로 작성하세요"""

        response = llm_chat(
            model=OLLAMA_MODEL,
            messages=[{"role": "user", "content": prompt}],
            options={"temperature": 0.2},
        )
        return response["message"]["content"].strip()

    except Exception as e:
        logger.warning(f"AI 서술 생성 스킵 ({topic}): {e}")
        return ""


def _add_narrative(doc: Document, topic: str, stats_text: str):
    """AI 서술 문단을 문서에 추가합니다. 생성 실패 시 조용히 건너뜁니다."""
    text = _generate_narrative(topic, stats_text)
    if not text:
        return
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(10)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)

# ─────────────────────────────────────────
# 색상 설정
# ─────────────────────────────────────────
PART_COLORS_RGB = {
    "안전계획": RGBColor(0x21, 0x96, 0xF3),
    "안전보건": RGBColor(0x4C, 0xAF, 0x50),
    "재난안전": RGBColor(0xFF, 0x98, 0x00),
}
RISK_COLORS_RGB = {
    "상": RGBColor(0xF4, 0x43, 0x36),
    "중": RGBColor(0xFF, 0xC1, 0x07),
    "하": RGBColor(0x4C, 0xAF, 0x50),
}
PART_COLORS_HEX = {
    "안전계획": "#2196F3",
    "안전보건": "#4CAF50",
    "재난안전": "#FF9800",
}
RISK_COLORS_HEX = {
    "상": "#F44336",
    "중": "#FFC107",
    "하": "#4CAF50",
}


# ─────────────────────────────────────────
# 한글 폰트 설정
# ─────────────────────────────────────────
def _set_korean_font():
    candidates = [
        "C:/Windows/Fonts/malgun.ttf",
        "C:/Windows/Fonts/gulim.ttc",
        "C:/Windows/Fonts/batang.ttc",
    ]
    for path in candidates:
        if Path(path).exists():
            font = fm.FontProperties(fname=path)
            matplotlib.rc("font", family=font.get_name())
            matplotlib.rcParams["axes.unicode_minus"] = False
            return


# ─────────────────────────────────────────
# 메인 보고서 생성 함수
# ─────────────────────────────────────────
def build_report(df: pd.DataFrame) -> Path:
    _set_korean_font()
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    # ★ 수정 (2026-07): 연도 하드코딩 제거 — 데이터의 최신 연도 사용,
    # 연도 데이터가 없으면 현재 연도로 대체
    report_year = _get_report_year(df)

    # 조치 이행상태 판정 컬럼 추가 (제6장 이행실태 분석에 사용)
    df = add_action_status(df)

    doc = Document()
    _setup_document(doc)

    _add_cover(doc, report_year)
    _add_toc(doc)
    _add_chapter_overview(doc, df)
    _add_chapter_classification(doc, df)
    _add_chapter_part_analysis(doc, df)
    _add_chapter_dept_analysis(doc, df)      # 신규: 제4장 부서별 분석
    _add_chapter_risk_analysis(doc, df)
    _add_chapter_action_analysis(doc, df)    # 신규: 제6장 조치 이행실태
    _add_chapter_repeat(doc, df)
    _add_chapter_recommendation(doc, df)

    filename = (
        f"{report_year}년_자체종합안전심사_AI분석보고서_"
        f"{datetime.now().strftime('%Y%m%d')}.docx"
    )
    save_path = OUTPUT_DIR / filename
    doc.save(save_path)
    logger.info(f"보고서 저장 완료: {save_path}")
    return save_path


def _get_report_year(df: pd.DataFrame) -> int:
    """보고서 기준 연도 — 데이터의 최신 연도, 없으면 현재 연도"""
    try:
        if "year" in df.columns:
            years = pd.to_numeric(df["year"], errors="coerce").dropna()
            if not years.empty:
                return int(years.max())
    except Exception:
        pass
    return datetime.now().year


# ─────────────────────────────────────────
# 문서 기본 설정
# ─────────────────────────────────────────
def _setup_document(doc: Document):
    section = doc.sections[0]
    section.page_width    = Inches(8.27)
    section.page_height   = Inches(11.69)
    section.left_margin   = Inches(1.0)
    section.right_margin  = Inches(1.0)
    section.top_margin    = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    style = doc.styles["Normal"]
    style.font.name = "맑은 고딕"
    style.font.size = Pt(10)


# ─────────────────────────────────────────
# 표지
# ─────────────────────────────────────────
def _add_cover(doc: Document, report_year: int):
    for _ in range(6):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"{report_year}년 자체종합안전심사")
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1A, 0x23, 0x7E)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run("AI 분석 결과 보고서")
    run2.font.size = Pt(24)
    run2.font.bold = True
    run2.font.color.rgb = RGBColor(0x1A, 0x23, 0x7E)

    doc.add_paragraph()

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = p3.add_run("철도안전관리체계 기반 AI 자동 분류 및 리스크 분석")
    run3.font.size = Pt(13)
    run3.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    for _ in range(5):
        doc.add_paragraph()

    info_data = [
        ("작  성  부  서", "안전계획팀"),
        ("작  성  일  자", datetime.now().strftime("%Y년 %m월 %d일")),
        ("분  석  모  델", "llama3.1:8b (로컬 AI)"),
        ("분  석  방  법", "RAG 기반 자동 분류"),
    ]
    for label, value in info_data:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_l = p.add_run(f"{label}  :  ")
        run_l.font.size = Pt(11)
        run_l.font.bold = True
        run_v = p.add_run(value)
        run_v.font.size = Pt(11)

    doc.add_page_break()


# ─────────────────────────────────────────
# 목차
# ─────────────────────────────────────────
def _add_toc(doc: Document):
    _add_heading(doc, "목  차", level=1)

    toc_items = [
        ("제1장", "분석 개요"),
        ("제2장", "AI 분류 결과 요약"),
        ("제3장", "파트별 세부 분석"),
        ("제4장", "부서별 분석"),
        ("제5장", "리스크 등급 분석"),
        ("제6장", "조치 이행실태 분석"),
        ("제7장", "반복 지적사항 분석"),
        ("제8장", "집중 점검 및 컨설팅 권고"),
    ]
    for num, title in toc_items:
        p = doc.add_paragraph()
        run = p.add_run(f"  {num}. {title}")
        run.font.size = Pt(11)

    doc.add_page_break()


# ─────────────────────────────────────────
# 1장: 분석 개요
# ─────────────────────────────────────────
def _add_chapter_overview(doc: Document, df: pd.DataFrame):
    _add_heading(doc, "제1장  분석 개요", level=1)

    _add_heading(doc, "1.1 분석 배경", level=2)
    doc.add_paragraph(
        "본 보고서는 대구도시철도공사(DTRO) 자체종합안전심사 이력 데이터를 "
        "AI 기반으로 자동 분류·분석하여 올해 심사의 집중 점검 방향을 도출하기 "
        "위해 작성되었습니다. llama3.1:8b 로컬 AI 모델과 RAG(Retrieval-Augmented "
        "Generation) 기법을 활용하여 과거 지적사항을 3개 파트로 분류하고 "
        "리스크 등급을 자동 판정하였습니다."
    )

    _add_heading(doc, "1.2 분석 범위", level=2)

    total         = len(df)
    ai_classified = df["ai_part"].notna().sum() if "ai_part" in df.columns else 0
    years         = sorted(df["year"].dropna().unique()) if "year" in df.columns else []

    table = doc.add_table(rows=5, cols=2)
    table.style = "Table Grid"

    rows_data = [
        ("분석 대상 연도",    ", ".join(str(int(y)) for y in years) if years else "2025"),
        ("총 지적사항 건수",  f"{total:,}건"),
        ("AI 분류 완료 건수", f"{ai_classified:,}건"),
        ("분류 파트",         "안전계획 / 안전보건 / 재난안전"),
        ("분석 도구",         "llama3.1:8b + nomic-embed-text (RAG)"),
    ]

    for i, (label, value) in enumerate(rows_data):
        row = table.rows[i]
        row.cells[0].text = label
        row.cells[1].text = value
        _set_cell_bg(row.cells[0], "E3F2FD")

    # 정렬 + 폭: [항목명 3.5cm, 내용 12.5cm]
    _format_table(table, [3.5, 12.5], header_row=False)

    # 항목명 Bold
    for row in table.rows:
        for para in row.cells[0].paragraphs:
            for run in para.runs:
                run.font.bold = True

    doc.add_paragraph()
    doc.add_page_break()


# ─────────────────────────────────────────
# 2장: AI 분류 결과 요약
# ─────────────────────────────────────────
def _add_chapter_classification(doc: Document, df: pd.DataFrame):
    _add_heading(doc, "제2장  AI 분류 결과 요약", level=1)

    if "ai_part" not in df.columns:
        doc.add_paragraph("AI 분류 데이터가 없습니다.")
        doc.add_page_break()
        return

    df_clean = df[
        df["ai_part"].notna() &
        (df["ai_part"] != "미분류")
    ].copy()

    part_counts = df_clean["ai_part"].value_counts()
    total       = len(df_clean)

    _add_heading(doc, "2.1 파트별 분류 현황", level=2)

    table = doc.add_table(rows=len(part_counts) + 2, cols=3)
    table.style = "Table Grid"

    # 헤더행
    headers = ["파트", "건수", "비율(%)"]
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = h
        _set_cell_bg(cell, "1565C0")

    # 데이터행
    for i, (part, count) in enumerate(part_counts.items(), 1):
        pct = count / total * 100
        row = table.rows[i]
        row.cells[0].text = part
        row.cells[1].text = f"{count:,}건"
        row.cells[2].text = f"{pct:.1f}%"
        color = {
            "안전계획": "E3F2FD",
            "안전보건": "E8F5E9",
            "재난안전": "FFF3E0"
        }.get(part, "FFFFFF")
        for cell in row.cells:
            _set_cell_bg(cell, color)

    # 합계행
    last_row = table.rows[len(part_counts) + 1]
    last_row.cells[0].text = "합  계"
    last_row.cells[1].text = f"{total:,}건"
    last_row.cells[2].text = "100.0%"
    for cell in last_row.cells:
        _set_cell_bg(cell, "E0E0E0")

    # 정렬 + 폭: [파트 4cm, 건수 4cm, 비율 4cm]
    _format_table(table, [4.0, 4.0, 4.0])

    # 헤더 Bold + 흰색
    _style_header_row(table.rows[0], white_text=True)
    # 합계 Bold
    for cell in last_row.cells:
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.bold = True

    doc.add_paragraph()

    # 차트 삽입
    chart_path = _create_part_chart(part_counts, total)
    if chart_path:
        doc.add_picture(str(chart_path), width=Inches(5.5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # AI 분석 서술 (신규) — 표의 수치를 근거로 한 해설 문단
    _add_heading(doc, "2.2 분석 종합", level=2)
    stats_text = "\n".join(
        f"- {part}: {count}건 ({count / total * 100:.1f}%)"
        for part, count in part_counts.items()
    ) + f"\n- 전체: {total}건"
    _add_narrative(doc, "파트별 지적사항 분포의 의미와 시사점", stats_text)

    doc.add_page_break()


# ─────────────────────────────────────────
# 3장: 파트별 세부 분석
# ─────────────────────────────────────────
def _add_chapter_part_analysis(doc: Document, df: pd.DataFrame):
    _add_heading(doc, "제3장  파트별 세부 분석", level=1)

    if "ai_part" not in df.columns:
        doc.add_paragraph("AI 분류 데이터가 없습니다.")
        doc.add_page_break()
        return

    parts = ["안전계획", "안전보건", "재난안전"]
    icons = {"안전계획": "🚇", "안전보건": "🏥", "재난안전": "🌪️"}

    for i, part in enumerate(parts, 1):
        df_part = df[df["ai_part"] == part].copy()
        if df_part.empty:
            continue

        _add_heading(doc, f"3.{i} {icons[part]} {part} 파트", level=2)
        doc.add_paragraph(
            f"{part} 파트에서는 총 {len(df_part)}건의 지적사항이 확인되었습니다."
        )

        if "audit_type" in df_part.columns:
            audit_counts = df_part["audit_type"].value_counts()
            p = doc.add_paragraph()
            for audit, cnt in audit_counts.items():
                p.add_run(f"  • {audit}: {cnt}건\n")

        _add_heading(doc, "주요 지적사항 TOP 5", level=3)

        top5 = (
            df_part.groupby("title")
            .size()
            .sort_values(ascending=False)
            .head(5)
            .reset_index()
        )
        top5.columns = ["지적사항", "건수"]

        if not top5.empty:
            table = doc.add_table(rows=len(top5) + 1, cols=3)
            table.style = "Table Grid"

            for j, h in enumerate(["순위", "지적사항 제목", "건수"]):
                cell = table.rows[0].cells[j]
                cell.text = h
                _set_cell_bg(cell, "1565C0")

            for idx, row_data in top5.iterrows():
                row = table.rows[idx + 1]
                row.cells[0].text = str(idx + 1)
                row.cells[1].text = str(row_data["지적사항"])
                row.cells[2].text = f"{row_data['건수']}건"

            # 정렬 + 폭: [순위 1.5cm, 지적사항 12cm, 건수 2.5cm]
            _format_table(table, [1.5, 12.0, 2.5])
            _style_header_row(table.rows[0], white_text=True)

        doc.add_paragraph()

    doc.add_page_break()


# ─────────────────────────────────────────
# 4장: 부서별 분석 (신규, 2026-07)
# ─────────────────────────────────────────
def _add_chapter_dept_analysis(doc: Document, df: pd.DataFrame):
    _add_heading(doc, "제4장  부서별 분석", level=1)

    if "department" not in df.columns:
        doc.add_paragraph("부서 데이터가 없습니다.")
        doc.add_page_break()
        return

    df_clean = df[
        df["department"].notna()
        & ~df["department"].isin(["미기재", "nan", ""])
    ]

    _add_heading(doc, "4.1 부서별 지적 현황 TOP 10", level=2)

    dept_counts = df_clean["department"].value_counts().head(10)
    table = doc.add_table(rows=len(dept_counts) + 1, cols=4)
    table.style = "Table Grid"

    for j, h in enumerate(["순위", "부서명", "지적건수", "리스크 상"]):
        cell = table.rows[0].cells[j]
        cell.text = h
        _set_cell_bg(cell, "1565C0")

    for i, (dept, cnt) in enumerate(dept_counts.items(), 1):
        dept_df = df_clean[df_clean["department"] == dept]
        high = len(dept_df[dept_df["ai_risk"] == "상"]) if "ai_risk" in dept_df.columns else 0
        row = table.rows[i]
        row.cells[0].text = str(i)
        row.cells[1].text = str(dept)
        row.cells[2].text = f"{cnt}건"
        row.cells[3].text = f"{high}건"
        if high > 0:
            _set_cell_bg(row.cells[3], "FFCDD2")

    _format_table(table, [1.5, 6.0, 3.0, 3.0])
    _style_header_row(table.rows[0], white_text=True)

    # AI 분석 서술
    _add_heading(doc, "4.2 분석 종합", level=2)
    stats_text = "\n".join(
        f"- {dept}: {cnt}건"
        for dept, cnt in dept_counts.items()
    )
    _add_narrative(doc, "부서별 지적 집중 현황과 관리 방향", stats_text)

    doc.add_paragraph()
    doc.add_page_break()


# ─────────────────────────────────────────
# 6장: 조치 이행실태 분석 (신규, 2026-07)
# ─────────────────────────────────────────
def _add_chapter_action_analysis(doc: Document, df: pd.DataFrame):
    _add_heading(doc, "제6장  조치 이행실태 분석", level=1)

    doc.add_paragraph(
        "추진실적 기록을 자동 판정하여 지적사항의 조치 이행 수준을 분석하였다. "
        "'형식적'은 구체적 내용 없이 완료 처리된 건으로, 실질 이행 여부의 확인이 필요하다."
    )

    if "action_status" not in df.columns:
        doc.add_paragraph("추진실적 데이터가 없습니다.")
        doc.add_page_break()
        return

    # ── 6.1 전체 이행 현황
    _add_heading(doc, "6.1 전체 이행 현황", level=2)
    s = action_summary(df)

    table = doc.add_table(rows=2, cols=6)
    table.style = "Table Grid"
    headers = ["전체", "완료", "형식적", "진행중", "미확인", "이행률"]
    values  = [
        f"{s['total']}건", f"{s['완료']}건", f"{s['형식적']}건",
        f"{s['진행중']}건", f"{s['미확인']}건", f"{s['이행률']}%",
    ]
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = h
        _set_cell_bg(cell, "1565C0")
    for j, v in enumerate(values):
        table.rows[1].cells[j].text = v
    _format_table(table, [2.5, 2.5, 2.5, 2.5, 2.5, 2.5])
    _style_header_row(table.rows[0], white_text=True)

    doc.add_paragraph()

    # ── 6.2 부서별 이행률 (낮은 순)
    _add_heading(doc, "6.2 부서별 이행률 (낮은 순)", level=2)
    dept_rate = action_rate_by_group(df, "department")

    if not dept_rate.empty:
        show = dept_rate.head(10)
        table2 = doc.add_table(rows=len(show) + 1, cols=5)
        table2.style = "Table Grid"

        for j, h in enumerate(["부서명", "지적건수", "완료", "형식적/미확인", "이행률"]):
            cell = table2.rows[0].cells[j]
            cell.text = h
            _set_cell_bg(cell, "1565C0")

        for i, (_, r) in enumerate(show.iterrows(), 1):
            row = table2.rows[i]
            row.cells[0].text = str(r["department"])
            row.cells[1].text = f"{r['지적건수']}건"
            row.cells[2].text = f"{r['완료']}건"
            row.cells[3].text = f"{r['형식적'] + r['미확인']}건"
            row.cells[4].text = f"{r['이행률(%)']}%"
            # 이행률 60% 이하 강조
            if r["이행률(%)"] <= 60:
                for cell in row.cells:
                    _set_cell_bg(cell, "FFCDD2")

        _format_table(table2, [5.0, 2.5, 2.0, 3.0, 2.5])
        _style_header_row(table2.rows[0], white_text=True)

    doc.add_paragraph()

    # ── 6.3 컨설팅 우선 대상 부서
    _add_heading(doc, "6.3 컨설팅 우선 대상 부서", level=2)
    doc.add_paragraph(
        "지적 3건 이상이면서 이행률 60% 이하인 부서로, "
        "지적 감소를 위한 컨설팅이 우선적으로 필요한 대상이다."
    )
    targets = consulting_targets(df)
    if targets.empty:
        doc.add_paragraph("해당 기준에 부합하는 부서가 없습니다.")
    else:
        for _, r in targets.iterrows():
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(
                f"{r['department']} — 지적 {r['지적건수']}건, "
                f"이행률 {r['이행률(%)']}% "
                f"(형식적 {r['형식적']}건, 미확인 {r['미확인']}건)"
            )

    # AI 분석 서술
    _add_heading(doc, "6.4 분석 종합", level=2)
    stats_text = (
        f"- 전체 {s['total']}건 중 완료 {s['완료']}건 (이행률 {s['이행률']}%)\n"
        f"- 형식적 완료 {s['형식적']}건, 진행중 {s['진행중']}건, 미확인 {s['미확인']}건\n"
    )
    if not dept_rate.empty:
        low3 = dept_rate.head(3)
        stats_text += "\n".join(
            f"- 이행률 하위 부서: {r['department']} {r['이행률(%)']}% (지적 {r['지적건수']}건)"
            for _, r in low3.iterrows()
        )
    _add_narrative(doc, "조치 이행실태의 수준 진단과 개선 방향", stats_text)

    doc.add_paragraph()
    doc.add_page_break()


# ─────────────────────────────────────────
# 5장: 리스크 등급 분석
# ─────────────────────────────────────────
def _add_chapter_risk_analysis(doc: Document, df: pd.DataFrame):
    _add_heading(doc, "제5장  리스크 등급 분석", level=1)

    if "ai_risk" not in df.columns:
        doc.add_paragraph("리스크 데이터가 없습니다.")
        doc.add_page_break()
        return

    df_clean = df[df["ai_part"].notna() & (df["ai_part"] != "미분류")]

    _add_heading(doc, "5.1 전체 리스크 등급 분포", level=2)
    risk_counts = df_clean["ai_risk"].value_counts()
    total_r     = len(df_clean)

    table = doc.add_table(rows=len(risk_counts) + 1, cols=3)
    table.style = "Table Grid"

    for j, h in enumerate(["리스크 등급", "건수", "비율"]):
        cell = table.rows[0].cells[j]
        cell.text = h
        _set_cell_bg(cell, "1565C0")

    for i, (risk, cnt) in enumerate(risk_counts.items(), 1):
        row = table.rows[i]
        row.cells[0].text = f"리스크 {risk}"
        row.cells[1].text = f"{cnt}건"
        row.cells[2].text = f"{cnt / total_r * 100:.1f}%"
        bg = {"상": "FFCDD2", "중": "FFF9C4", "하": "C8E6C9"}.get(risk, "FFFFFF")
        for cell in row.cells:
            _set_cell_bg(cell, bg)

    # 정렬 + 폭: [등급 4cm, 건수 4cm, 비율 4cm]
    _format_table(table, [4.0, 4.0, 4.0])
    _style_header_row(table.rows[0], white_text=True)

    doc.add_paragraph()

    _add_heading(doc, "5.2 리스크 '상' 항목 목록 (집중 관리 필요)", level=2)
    df_high = df_clean[df_clean["ai_risk"] == "상"].copy()

    if not df_high.empty:
        doc.add_paragraph(
            f"리스크 '상' 등급 항목은 총 {len(df_high)}건으로 즉각적인 조치가 필요합니다."
        )
        cols  = ["ai_part", "audit_type", "title", "department"]
        avail = [c for c in cols if c in df_high.columns]
        show  = df_high[avail].head(15)

        col_map = {
            "ai_part":    "파트",
            "audit_type": "심사구분",
            "title":      "지적사항",
            "department": "담당부서"
        }

        table2 = doc.add_table(rows=len(show) + 1, cols=len(avail))
        table2.style = "Table Grid"

        for j, c in enumerate(avail):
            cell = table2.rows[0].cells[j]
            cell.text = col_map.get(c, c)
            _set_cell_bg(cell, "1565C0")

        for i, (_, row_data) in enumerate(show.iterrows(), 1):
            for j, c in enumerate(avail):
                table2.rows[i].cells[j].text = str(row_data[c])
                _set_cell_bg(table2.rows[i].cells[j], "FFCDD2")

        # 정렬 + 폭: [파트 2.5cm, 심사구분 2.5cm, 지적사항 8cm, 담당부서 3cm]
        _format_table(table2, [2.5, 2.5, 8.0, 3.0])
        _style_header_row(table2.rows[0], white_text=True)

    doc.add_paragraph()
    doc.add_page_break()


# ─────────────────────────────────────────
# 5장: 반복 지적사항
# ─────────────────────────────────────────
def _add_chapter_repeat(doc: Document, df: pd.DataFrame):
    _add_heading(doc, "제7장  반복 지적사항 분석", level=1)

    doc.add_paragraph(
        "반복적으로 지적되는 사항은 구조적 문제일 가능성이 높으므로 "
        "근본 원인 분석 및 체계적 개선이 필요합니다."
    )

    if "title" not in df.columns:
        doc.add_page_break()
        return

    repeat = (
        df.groupby(["title", "ai_part"])
        .size()
        .reset_index(name="반복횟수")
        .sort_values("반복횟수", ascending=False)
        .head(10)
    )

    if repeat.empty:
        doc.add_paragraph("반복 지적사항이 없습니다.")
        doc.add_page_break()
        return

    table = doc.add_table(rows=len(repeat) + 1, cols=3)
    table.style = "Table Grid"

    for j, h in enumerate(["지적사항", "파트", "반복횟수"]):
        cell = table.rows[0].cells[j]
        cell.text = h
        _set_cell_bg(cell, "1565C0")

    for i, (_, row_data) in enumerate(repeat.iterrows(), 1):
        row = table.rows[i]
        row.cells[0].text = str(row_data["title"])
        row.cells[1].text = str(row_data["ai_part"])
        row.cells[2].text = f"{row_data['반복횟수']}회"

        cnt = row_data["반복횟수"]
        bg  = "FFCDD2" if cnt >= 3 else "FFF9C4" if cnt >= 2 else "FFFFFF"
        for cell in row.cells:
            _set_cell_bg(cell, bg)

    # 정렬 + 폭: [지적사항 10cm, 파트 3cm, 반복횟수 3cm]
    _format_table(table, [10.0, 3.0, 3.0])
    _style_header_row(table.rows[0], white_text=True)

    doc.add_paragraph()
    doc.add_page_break()


# ─────────────────────────────────────────
# 6장: 집중 점검 권고사항
# ─────────────────────────────────────────
def _add_chapter_recommendation(doc: Document, df: pd.DataFrame):
    _add_heading(doc, "제8장  집중 점검 및 컨설팅 권고", level=1)

    doc.add_paragraph(
        "AI 분석 결과를 바탕으로 올해 자체종합안전심사에서 "
        "집중적으로 점검해야 할 항목을 파트별로 제시합니다."
    )

    if "ai_part" not in df.columns:
        doc.add_page_break()
        return

    parts_info = {
        "안전계획": {"icon": "🚇", "desc": "철도사고·준사고·운행장애 예방 관련"},
        "안전보건": {"icon": "🏥", "desc": "산업안전보건법 준수 및 근로자 보건 관련"},
        "재난안전": {"icon": "🌪️", "desc": "자연재난 대응 및 비상훈련 관련"},
    }

    for i, (part, info) in enumerate(parts_info.items(), 1):
        df_part = df[df["ai_part"] == part]
        if df_part.empty:
            continue

        _add_heading(doc, f"8.{i} {info['icon']} {part} 파트 집중 점검 항목", level=2)
        doc.add_paragraph(f"● 분야: {info['desc']}")
        doc.add_paragraph(f"● 해당 건수: {len(df_part)}건")

        if "ai_risk" in df_part.columns:
            high_risk = df_part[df_part["ai_risk"] == "상"]
            if not high_risk.empty:
                doc.add_paragraph(
                    f"● 리스크 '상' 항목 ({len(high_risk)}건) — 즉시 조치 필요"
                )
                for _, row in high_risk.head(5).iterrows():
                    p = doc.add_paragraph(style="List Bullet")
                    p.add_run(str(row.get("title", "")))

        repeat_items = (
            df_part.groupby("title")
            .size()
            .sort_values(ascending=False)
            .head(3)
        )
        doc.add_paragraph("● 반복 지적 항목 (우선 점검)")
        for title, cnt in repeat_items.items():
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(f"{title} ({cnt}회 반복)")

        # AI 컨설팅 서술 (신규) — 통계 근거 기반 개선 컨설팅 방향 제시
        doc.add_paragraph("● AI 컨설팅 의견")
        act = action_summary(df_part) if "action_status" in df_part.columns else None
        stats_text = f"- {part} 파트 지적 {len(df_part)}건\n"
        if "ai_risk" in df_part.columns:
            stats_text += f"- 리스크 상 {len(df_part[df_part['ai_risk'] == '상'])}건\n"
        stats_text += "\n".join(
            f"- 반복 지적: {title} ({cnt}회)"
            for title, cnt in repeat_items.items()
        )
        if act:
            stats_text += f"\n- 조치 이행률 {act['이행률']}% (형식적 {act['형식적']}건, 미확인 {act['미확인']}건)"
        _add_narrative(
            doc,
            f"{part} 파트의 반복 지적 원인 추정과 지적 감소를 위한 컨설팅 방향 (컨설팅 포인트 2~3개 포함)",
            stats_text
        )

        doc.add_paragraph()

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        "본 보고서는 AI 자동 분석 결과이며, "
        "최종 심사 계획 수립 시 실무자 검토를 반드시 병행하시기 바랍니다."
    )
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)


# ─────────────────────────────────────────
# 차트 생성 (워드 삽입용 PNG)
# ─────────────────────────────────────────
def _create_part_chart(part_counts, total) -> Path | None:
    try:
        fig, axes = plt.subplots(1, 2, figsize=(10, 4))

        parts  = list(part_counts.index)
        counts = list(part_counts.values)
        colors = [PART_COLORS_HEX.get(p, "#999") for p in parts]

        # 막대 차트
        bars = axes[0].bar(parts, counts, color=colors, edgecolor="white", linewidth=2)
        for bar, count in zip(bars, counts):
            pct = count / total * 100
            axes[0].text(
                bar.get_x() + bar.get_width() / 2,
                bar.get_height() + 0.5,
                f"{count}건\n({pct:.1f}%)",
                ha="center", va="bottom", fontsize=9
            )
        axes[0].set_title("파트별 분류 현황", fontsize=13)
        axes[0].set_ylabel("건수")
        axes[0].set_ylim(0, max(counts) * 1.35)
        axes[0].grid(axis="y", alpha=0.3)

        # 파이 차트
        axes[1].pie(
            counts,
            labels=parts,
            colors=colors,
            autopct="%1.1f%%",
            startangle=90,
            wedgeprops=dict(edgecolor="white", linewidth=2)
        )
        axes[1].set_title("파트별 비율", fontsize=13)

        plt.tight_layout()
        chart_path = OUTPUT_DIR / "temp_chart.png"
        plt.savefig(chart_path, dpi=150, bbox_inches="tight")
        plt.close()
        return chart_path

    except Exception as e:
        logger.error(f"차트 생성 오류: {e}")
        return None


# ─────────────────────────────────────────
# 유틸리티 — 제목
# ─────────────────────────────────────────
def _add_heading(doc: Document, text: str, level: int):
    p   = doc.add_paragraph()
    run = p.add_run(text)

    if level == 1:
        run.font.size      = Pt(16)
        run.font.bold      = True
        run.font.color.rgb = RGBColor(0x1A, 0x23, 0x7E)
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after  = Pt(6)
    elif level == 2:
        run.font.size      = Pt(13)
        run.font.bold      = True
        run.font.color.rgb = RGBColor(0x15, 0x65, 0xC0)
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after  = Pt(4)
    elif level == 3:
        run.font.size      = Pt(11)
        run.font.bold      = True
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after  = Pt(3)


# ─────────────────────────────────────────
# 유틸리티 — 셀 배경색
# ─────────────────────────────────────────
def _set_cell_bg(cell, hex_color: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


# ─────────────────────────────────────────
# 유틸리티 — 셀 가운데 정렬
# ─────────────────────────────────────────
def _set_cell_center(cell):
    """셀 내용 가로 + 세로 가운데 정렬"""
    # 세로 가운데
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # 가로 가운데
    for para in cell.paragraphs:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in para.runs:
            run.font.name = "맑은 고딕"


# ─────────────────────────────────────────
# 유틸리티 — 컬럼 폭 설정
# ─────────────────────────────────────────
def _set_col_widths(table, col_widths_cm: list):
    """컬럼 폭 설정 (cm 단위, 1cm = 567 twips)"""
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            if i >= len(col_widths_cm):
                continue
            tc   = cell._tc
            tcPr = tc.get_or_add_tcPr()

            # 기존 tcW 제거
            for child in tcPr.findall(qn("w:tcW")):
                tcPr.remove(child)

            tcW = OxmlElement("w:tcW")
            tcW.set(qn("w:w"),    str(int(col_widths_cm[i] * 567)))
            tcW.set(qn("w:type"), "dxa")
            tcPr.append(tcW)


# ─────────────────────────────────────────
# 유틸리티 — 표 전체 포맷 (정렬 + 폭)
# ─────────────────────────────────────────
def _format_table(table, col_widths_cm: list, header_row: bool = True):
    """
    표 전체에 가운데 정렬 + 컬럼 폭 적용
    :param col_widths_cm: 각 컬럼 폭 (cm)
    :param header_row: True면 0번 행은 헤더로 처리
    """
    _set_col_widths(table, col_widths_cm)

    for i, row in enumerate(table.rows):
        row.height = Pt(20)
        for cell in row.cells:
            _set_cell_center(cell)


# ─────────────────────────────────────────
# 유틸리티 — 헤더행 스타일
# ─────────────────────────────────────────
def _style_header_row(row, white_text: bool = True):
    """헤더행 Bold + 선택적 흰색 글자"""
    for cell in row.cells:
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.bold = True
                if white_text:
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)